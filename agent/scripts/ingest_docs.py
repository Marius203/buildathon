"""
Chunk EC_Documents/ (docx + xlsx + txt) and write to data/chunks.json.
Does NOT embed — run embed_chunks.py afterwards.

Run from the agent/ directory:
    python -m scripts.ingest_docs
    python -m scripts.embed_chunks
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

import docx
import openpyxl

sys.path.insert(0, str(Path(__file__).parent.parent))

DOCS_DIR = Path(__file__).parent.parent / "EC_Documents"
DUMP_PATH = Path(__file__).parent.parent / "data" / "chunks.json"

HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "heading 4"}
_DAYS = {"MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"}
_SKIP_PHRASES = {
    "USE THESE JUST AS AN EXAMPLE FOR THE TONE OF VOICE AND DON'T",
    "***Question examples start***",
}

_ROMAN = re.compile(r"^[IVX]+\.\s")


def make_id(source: str, idx: int) -> str:
    return hashlib.md5(f"{source}::{idx}".encode()).hexdigest()


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _build_record(source: str, idx: int, chunk: dict) -> dict:
    return {
        "id": make_id(source, idx),
        "text": chunk["text"],
        "lang": "en",
        "section": chunk.get("section", ""),
        "section_title": chunk.get("section_title", ""),
        "source": source,
        "created_at": _now(),
    }


# ── XLSX ──────────────────────────────────────────────────────────────────────

def chunk_xlsx(path: Path) -> list[dict]:
    wb = openpyxl.load_workbook(str(path), data_only=True)
    ws = wb["Final lineup table"]
    chunks: list[dict] = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        artist, time_place, description, genre = (row[i] for i in range(4))
        if not artist:
            continue
        parts = [f"Artist: {artist}"]
        if time_place:
            parts.append(f"When/Where: {time_place}")
        if genre and str(genre).strip() not in ("#N/A", ""):
            parts.append(f"Genre: {genre}")
        if description:
            parts.append(f"Description: {description}")
        chunks.append({
            "text": "\n".join(parts),
            "section_title": str(artist),
            "section": "lineup",
        })

    return chunks


# ── DOCX: heading-based (Rules, Tickets) ──────────────────────────────────────

def chunk_docx_headings(path: Path) -> list[dict]:
    doc_obj = docx.Document(str(path))
    chunks: list[dict] = []
    current_title = path.stem
    buf: list[str] = []

    def flush():
        text = "\n".join(buf).strip()
        if text:
            chunks.append({"text": text, "section_title": current_title, "section": current_title})
        buf.clear()

    for para in doc_obj.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        if not text:
            continue
        if style in HEADING_STYLES:
            flush()
            current_title = text
        else:
            buf.append(text)

    flush()
    return chunks


# ── DOCX: day-based (Daily Recommendations) ───────────────────────────────────

def chunk_docx_daily(path: Path) -> list[dict]:
    doc_obj = docx.Document(str(path))
    chunks: list[dict] = []
    current_title = "General Recommendations"
    buf: list[str] = []

    def flush():
        text = "\n".join(buf).strip()
        if text:
            chunks.append({"text": text, "section_title": current_title, "section": current_title})
        buf.clear()

    for para in doc_obj.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        if not text:
            continue
        if style in HEADING_STYLES:
            continue
        if text.upper() in _DAYS:
            flush()
            current_title = f"{text.capitalize()} Recommendations"
        else:
            buf.append(text)

    flush()
    return chunks


# ── DOCX: Q&A pairs ───────────────────────────────────────────────────────────

def chunk_docx_qa(path: Path) -> list[dict]:
    doc_obj = docx.Document(str(path))
    paras = [
        p.text.strip()
        for p in doc_obj.paragraphs
        if p.text.strip() and p.text.strip() not in _SKIP_PHRASES
    ]

    topic_indices: set[int] = set()
    for i, text in enumerate(paras):
        if (
            not text.startswith("-")
            and len(text) < 60
            and not text.endswith(".")
            and i + 1 < len(paras)
            and paras[i + 1].startswith("-")
        ):
            topic_indices.add(i)

    chunks: list[dict] = []
    current_section = "General"
    current_q: str | None = None
    current_a: list[str] = []

    def flush_qa():
        nonlocal current_q
        if current_q:
            answer = " ".join(current_a).strip()
            text = f"Q: {current_q}\nA: {answer}" if answer else f"Q: {current_q}"
            chunks.append({"text": text, "section_title": current_section, "section": current_section})
        current_q = None
        current_a.clear()

    for i, text in enumerate(paras):
        if i in topic_indices:
            current_section = text
        elif text.startswith("-"):
            flush_qa()
            current_q = text.lstrip("- ").strip()
        elif current_q is not None:
            current_a.append(text)

    flush_qa()
    return chunks


# ── DOCX: single chunk (Important Links) ─────────────────────────────────────

def chunk_docx_links(path: Path) -> list[dict]:
    doc_obj = docx.Document(str(path))
    lines = [p.text.strip() for p in doc_obj.paragraphs if p.text.strip()]
    return [{"text": "\n".join(lines), "section_title": "Important Links", "section": "links"}]


# ── DOCX router ───────────────────────────────────────────────────────────────

def chunk_docx(path: Path) -> list[dict]:
    stem = path.stem.lower()
    if "recommendation" in stem:
        return chunk_docx_daily(path)
    if "question" in stem:
        return chunk_docx_qa(path)
    if "link" in stem:
        return chunk_docx_links(path)
    return chunk_docx_headings(path)


# ── TXT ───────────────────────────────────────────────────────────────────────

def _is_txt_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if _ROMAN.match(stripped):
        return True
    words = stripped.split()
    if len(words) < 2:
        return False
    if stripped != stripped.upper():
        return False
    alpha_only = re.sub(r"[\s\-/&]", "", stripped)
    return alpha_only.isalpha()


def chunk_txt(path: Path) -> list[dict]:
    chunks: list[dict] = []
    current_title = path.stem
    buf: list[str] = []

    def flush():
        text = "\n".join(buf).strip()
        if text:
            chunks.append({"text": text, "section_title": current_title, "section": current_title})
        buf.clear()

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if _is_txt_heading(line):
            flush()
            current_title = line
        else:
            buf.append(line)

    flush()
    return chunks


# ── Main ──────────────────────────────────────────────────────────────────────

def ingest_all():
    all_chunks: list[tuple[str, dict]] = []

    for path in sorted(DOCS_DIR.iterdir()):
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            continue
        if suffix == ".docx":
            chunks = chunk_docx(path)
        elif suffix in (".xlsx", ".xls"):
            chunks = chunk_xlsx(path)
        elif suffix == ".txt":
            chunks = chunk_txt(path)
        else:
            continue
        print(f"{path.name}: {len(chunks)} chunks", flush=True)
        for c in chunks:
            all_chunks.append((path.name, c))

    records = [_build_record(src, i, c) for i, (src, c) in enumerate(all_chunks)]

    DUMP_PATH.parent.mkdir(parents=True, exist_ok=True)
    DUMP_PATH.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nTotal: {len(records)} chunks → {DUMP_PATH}")


if __name__ == "__main__":
    ingest_all()
